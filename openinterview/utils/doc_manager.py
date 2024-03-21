import os
from datetime import datetime
from collections import defaultdict
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.shared import OxmlElement, qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING


class DocumentCreator:
    def __init__(self):
        self.doc = Document()

    def set_default_font(self, font_name, font_size):
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = font_name
        if font_size is not None:
            font.size = font_size

    @staticmethod
    def set_cell_background_black(cell):
        shading_elm = parse_xml(r'<w:shd {} w:fill="000000"/>'.format(nsdecls("w")))
        cell._tc.get_or_add_tcPr().append(shading_elm)

    @staticmethod
    def set_cell_margins(cell, top, bottom, left, right):
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_mar = OxmlElement("w:tcMar")
        for name, value in [
            ("top", top),
            ("bottom", bottom),
            ("left", left),
            ("right", right),
        ]:
            mar = OxmlElement(f"w:{name}")
            mar.set(qn("w:w"), str(value * 20))
            mar.set(qn("w:type"), "dxa")
            tc_mar.append(mar)
        tc_pr.append(tc_mar)

    @staticmethod
    def set_table_borders_to_light_gray(table):
        for row in table.rows:
            for cell in row.cells:
                for key in ["top", "left", "bottom", "right"]:
                    tag = f"{key}OrTblGrid"
                    border = OxmlElement(f"w:tcBorders")
                    border_sub = OxmlElement(f"w:{tag}")
                    border_sub.set(qn("w:val"), "single")
                    border_sub.set(qn("w:sz"), "7")
                    border_sub.set(qn("w:color"), "D3D3D3")
                    border.append(border_sub)
                    cell._tc.get_or_add_tcPr().append(border)

    @staticmethod
    def set_cell_padding(cell, padding_pt=2):
        padding_twips = padding_pt * 20
        tc_pr = cell._tc.get_or_add_tcPr()
        for side in ["top", "left", "bottom", "right"]:
            tc_mar = OxmlElement(f"w:tcMar")
            side_mar = OxmlElement(f"w:{side}")
            side_mar.set(qn("w:w"), str(padding_twips))
            side_mar.set(qn("w:type"), "dxa")
            tc_mar.append(side_mar)
            tc_pr.append(tc_mar)

    @staticmethod
    def set_paragraph_spacing(
        paragraph, line_spacing=None, space_before=None, space_after=None
    ):
        if line_spacing is not None:
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            paragraph.paragraph_format.line_spacing = line_spacing
        if space_before is not None:
            paragraph.paragraph_format.space_before = Pt(space_before)
        if space_after is not None:
            paragraph.paragraph_format.space_after = Pt(space_after)

    def create_qa_document(self, qa_dict, font_name, font_size, save_dir):
        self.set_default_font(font_name, font_size)

        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.59)  # 15mm
            section.bottom_margin = Inches(0.59)  # 15mm
            section.left_margin = Inches(0.59)  # 15mm
            section.right_margin = Inches(0.59)  # 15mm

        # Project Information Section
        title = self.doc.add_paragraph("Open Interview")
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title.runs[0].font.size = Pt(20)

        install_cmd = self.doc.add_paragraph("$ pip install open-interview")
        install_cmd.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        install_cmd.runs[0].font.size = Pt(13)

        self.doc.add_paragraph()
        author_cmd = self.doc.add_paragraph(
            "Author: Minwoo(Daniel) Park\nGitHub Repository: https://github.com/dsdanielpark/open-interview"
        )
        author_cmd.runs[0].font.size = Pt(8)
        self.doc.add_paragraph()

        # Organize Q&A pairs by the identifier after the first underscore
        qa_groups = defaultdict(lambda: {"Q": None, "A": None})
        for key, value in qa_dict.items():
            try:
                prefix, identifier = key.split("_", 1)
                qa_groups[identifier][prefix] = value
            except:
                pass
        # Iterate through the grouped Q&A pairs to create tables in the document
        for identifier, qa_pair in qa_groups.items():
            table = self.doc.add_table(rows=2, cols=1, style="Table Grid")
            self.set_table_borders_to_light_gray(table)

            # Question Cell
            q_cell = table.cell(0, 0)
            q_cell.text = "Q: " + (
                qa_pair["Q"] if qa_pair["Q"] else "No question provided."
            )
            self.set_cell_background_black(q_cell)
            self.set_cell_margins(q_cell, 3, 3, 3, 3)
            self.set_cell_padding(q_cell, 4)
            q_paragraph = q_cell.paragraphs[0]
            q_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            self.set_paragraph_spacing(q_paragraph, line_spacing=1.5)
            q_run = q_paragraph.runs[0]
            q_run.font.color.rgb = RGBColor(255, 255, 255)
            q_run.font.size = Pt(12)
            q_run.font.bold = False

            # Answer Cell
            a_cell = table.cell(1, 0)
            a_cell.text = "A: " + (
                qa_pair["A"] if qa_pair["A"] else "No answer provided."
            )
            self.set_cell_margins(a_cell, 3, 3, 3, 3)
            self.set_cell_padding(a_cell, 4)
            a_paragraph = a_cell.paragraphs[0]
            a_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            self.set_paragraph_spacing(a_paragraph, line_spacing=1.5)
            a_run = a_paragraph.runs[0]
            a_run.font.size = Pt(12)
            a_run.font.bold = False

            self.doc.add_paragraph()  # Add an empty paragraph for spacing between Q&A groups

        os.makedirs(save_dir, exist_ok=True)
        save_path = (
            f"{save_dir}/OpenInterview_{datetime.now().strftime('%Y%m%d%H%M%S%f')}.docx"
        )
        self.doc.save(save_path)
